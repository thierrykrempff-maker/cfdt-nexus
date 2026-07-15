"""Local PDF/DOCX importer without OCR, network, AI, or business analysis."""
from __future__ import annotations
import argparse, hashlib, json, os, re, time, uuid
from collections import Counter, defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Callable
from automation.protection_sociale.audit_corpus import classify_path, sha256_file
from automation.protection_sociale.import_models import ImportedDocument, ExtractionResult, IMPORT_VERSION, SCHEMA_VERSION

SUPPORTED={".pdf",".docx"}; ABSOLUTE=re.compile(r"(?i)(?:[a-z]:[\\/][^\r\n\t<>\"|?*]+|/(?:home|users)/[^\r\n\t<>\"|?*]+)")
def now(): return datetime.now(timezone.utc).isoformat()
def stable_document_id(relative:str)->str: return str(uuid.uuid5(uuid.NAMESPACE_URL,f"cfdt-nexus:protection-sociale:import:{Path(relative).as_posix()}:{IMPORT_VERSION}"))
def duplicate_group_id(digest:str)->str: return str(uuid.uuid5(uuid.NAMESPACE_URL,f"cfdt-nexus:protection-sociale:duplicate:{digest}"))
def redact(value:str|None)->tuple[str|None,bool]:
    if value is None:return None,False
    clean,count=ABSOLUTE.subn("[ABSOLUTE_PATH_REDACTED]",value); return clean,bool(count)

def extract_pdf(path:Path)->ExtractionResult:
    from pypdf import PdfReader
    reader=PdfReader(path); parts=[]; warnings=[]
    for i,page in enumerate(reader.pages,1):
        text=page.extract_text() or ""
        if not text.strip(): warnings.append(f"page_without_text:{i}")
        parts.append(f"--- PAGE {i} ---\n{text.strip()}")
    if reader.pages and all(x.startswith("page_without_text") for x in warnings): warnings.append("probable_image_only_pdf_no_ocr")
    return ExtractionResult("\n\n".join(parts),"extracted_with_warnings" if warnings else "extracted","pypdf","page.extract_text (no OCR)",len(reader.pages),None,None,warnings)

def extract_docx(path:Path)->ExtractionResult:
    from docx import Document
    doc=Document(path); parts=[]
    for i,p in enumerate(doc.paragraphs,1):
        if p.text.strip(): parts.append(f"--- PARAGRAPH {i} ---\n{p.text}")
    for i,t in enumerate(doc.tables,1):
        parts.append(f"--- TABLE {i} ---")
        parts.extend("\t".join(c.text for c in row.cells) for row in t.rows)
    return ExtractionResult("\n\n".join(parts),"extracted","python-docx","paragraphs and tables",None,len(doc.paragraphs),len(doc.tables),[])
EXTRACTORS:dict[str,Callable[[Path],ExtractionResult]]={".pdf":extract_pdf,".docx":extract_docx}

def validate_paths(source:Path,output:Path)->tuple[Path,Path]:
    source=source.resolve(); output=output.resolve()
    if not source.is_dir(): raise FileNotFoundError("Source directory does not exist")
    if output==source or source in output.parents or "RAW_DOCUMENTS" in str(output).replace("\\","/").upper(): raise ValueError("Output must not be inside RAW_DOCUMENTS")
    if source.is_symlink() or output.is_symlink(): raise ValueError("Symbolic links are refused")
    return source,output

def discover(source:Path)->list[Path]:
    found=[]
    for current,dirs,files in os.walk(source,followlinks=False):
        dirs[:]=[d for d in dirs if not (Path(current)/d).is_symlink()]
        for name in files:
            p=Path(current)/name
            if not p.is_symlink() and p.suffix.casefold()!=".lnk": found.append(p)
    return sorted(found,key=lambda p:p.relative_to(source).as_posix().casefold())

def hash_inventory(files:list[Path])->tuple[dict[Path,str],dict[str,list[Path]]]:
    hashes={}; groups=defaultdict(list)
    for p in files:
        try: digest=sha256_file(p); hashes[p]=digest; groups[digest].append(p)
        except OSError: pass
    return hashes,groups

def build_record(path:Path,source:Path,digest:str,duplicates:dict[str,list[Path]])->ImportedDocument:
    relative=path.relative_to(source).as_posix(); stat=path.stat(); domain,category,subcategory=classify_path(relative); warnings=[]; error_code=error_message=None
    if stat.st_size==0: result=ExtractionResult(status="empty",extractor_name="none",extractor_method="size_check",warnings=["empty_source"])
    elif path.suffix.casefold() not in EXTRACTORS: result=ExtractionResult(status="unsupported",extractor_name="none",extractor_method="not_opened",warnings=["unsupported_extension"])
    else:
        try: result=EXTRACTORS[path.suffix.casefold()](path)
        except Exception as exc:
            result=ExtractionResult(status="failed",extractor_name=EXTRACTORS[path.suffix.casefold()].__name__,extractor_method="direct")
            error_code=type(exc).__name__; error_message=str(exc)[:500]
    safe_text,text_changed=redact(result.text_content); safe_error,error_changed=redact(error_message)
    if text_changed or error_changed: result.warnings.append("absolute_path_redacted"); result.status="extracted_with_warnings" if result.status=="extracted" else result.status
    duplicate=len(duplicates.get(digest,[]))>1
    return ImportedDocument(stable_document_id(relative),relative,path.name,path.suffix.casefold(),stat.st_size,digest,datetime.fromtimestamp(stat.st_mtime,timezone.utc).isoformat(),domain,category,subcategory,result.extractor_name,result.extractor_method,result.status,error_code,safe_error,safe_text or "",len(safe_text or ""),result.page_count,result.paragraph_count,result.table_count,duplicate_group_id(digest) if duplicate else None,duplicate,stat.st_size==0,result.warnings,now())

def write_json(path:Path,value)->None: path.parent.mkdir(parents=True,exist_ok=True); path.write_text(json.dumps(value,ensure_ascii=False,indent=2)+"\n",encoding="utf-8")
def balanced_sample(files:list[Path],count:int=4)->list[Path]:
    nonempty_pdf=[p for p in files if p.suffix.casefold()==".pdf" and p.stat().st_size>0]; nonempty_docx=[p for p in files if p.suffix.casefold()==".docx" and p.stat().st_size>0]; empty=[p for p in files if p.stat().st_size==0]
    selected=nonempty_pdf[:2]+nonempty_docx[:1]+empty[:1]
    for p in files:
        if len(selected)>=count: break
        if p not in selected:selected.append(p)
    return selected[:count]

def run_import(source:Path|str,output:Path|str,*,mode="dry-run",extensions:set[str]|None=None,domain=None,category=None,subfolder=None,limit=None,balanced=False,force=False)->dict:
    source,output=validate_paths(Path(source),Path(output)); all_files=discover(source); hashes,groups=hash_inventory(all_files); files=all_files
    if extensions: files=[p for p in files if p.suffix.casefold() in {x.casefold() if x.startswith('.') else '.'+x.casefold() for x in extensions}]
    if subfolder: files=[p for p in files if p.relative_to(source).as_posix().casefold().startswith(subfolder.replace('\\','/').casefold().rstrip('/')+'/')]
    if domain or category: files=[p for p in files if (not domain or classify_path(p.relative_to(source).as_posix())[0]==domain) and (not category or classify_path(p.relative_to(source).as_posix())[1]==category)]
    if balanced: files=balanced_sample(files)
    if limit is not None: files=files[:max(0,limit)]
    started=time.monotonic(); statuses=Counter(); documents=[]; errors=[]; processed=resumed=text_chars=pages=paragraphs=tables=0
    for p in files:
        relative=p.relative_to(source).as_posix(); digest=hashes.get(p)
        if digest is None: statuses['failed']+=1; errors.append({"source_relative_path":relative,"error_code":"HashReadError"}); continue
        doc_id=stable_document_id(relative); target=output/"documents"/(doc_id+".json")
        if mode=="import" and target.exists() and not force:
            try:
                previous=json.loads(target.read_text(encoding="utf-8"))
                if previous.get("source_sha256")==digest: resumed+=1; statuses[previous["extraction_status"]]+=1; text_chars+=previous.get("text_length",0); pages+=previous.get("page_count") or 0; paragraphs+=previous.get("paragraph_count") or 0; tables+=previous.get("table_count") or 0; documents.append({"document_id":doc_id,"source_relative_path":relative,"status":previous["extraction_status"],"resumed":True}); continue
            except (OSError,ValueError,KeyError): pass
        if mode=="dry-run":
            status="empty" if p.stat().st_size==0 else ("extractable" if p.suffix.casefold() in SUPPORTED else "unsupported"); statuses[status]+=1; documents.append({"document_id":doc_id,"source_relative_path":relative,"planned_status":status}); continue
        record=build_record(p,source,digest,groups); write_json(target,record.to_dict()); processed+=1; statuses[record.extraction_status]+=1; text_chars+=record.text_length; pages+=record.page_count or 0; paragraphs+=record.paragraph_count or 0; tables+=record.table_count or 0
        documents.append({"document_id":doc_id,"source_relative_path":relative,"status":record.extraction_status,"resumed":False})
        if record.extraction_error_code: errors.append({"document_id":doc_id,"source_relative_path":relative,"error_code":record.extraction_error_code,"error_message":record.extraction_error_message})
    duplicate_groups=[{"duplicate_group_id":duplicate_group_id(h),"source_sha256":h,"document_ids":[stable_document_id(p.relative_to(source).as_posix()) for p in paths],"copy_count":len(paths)} for h,paths in groups.items() if len(paths)>1]
    manifest={"schema_version":SCHEMA_VERSION,"mode":mode,"generated_at":now(),"examined_count":len(files),"processed_count":processed,"resumed_count":resumed,"status_counts":dict(sorted(statuses.items())),"text_length_total":text_chars,"page_count_total":pages,"paragraph_count_total":paragraphs,"table_count_total":tables,"duplicate_group_count":len(duplicate_groups),"error_count":len(errors),"duration_seconds":round(time.monotonic()-started,3),"documents":documents}
    if mode=="import":
        write_json(output/"manifests"/"import_manifest.json",manifest); write_json(output/"manifests"/"duplicate_report.json",{"duplicate_groups":duplicate_groups}); write_json(output/"logs"/"import_errors.json",{"errors":errors})
        summary=f"# Import Protection sociale — LOT 1A\n\nDocuments examinés : {len(files)}\n\nDocuments traités : {processed}\n\nCaractères extraits : {text_chars}\n\nAucun OCR, réseau, IA ou analyse métier.\n"; (output/"manifests"/"import_summary.md").write_text(summary,encoding="utf-8")
    return manifest

def main()->int:
    p=argparse.ArgumentParser(description=__doc__); p.add_argument("--source",type=Path,required=True); p.add_argument("--output",type=Path,required=True); p.add_argument("--mode",choices=("dry-run","import"),default="dry-run"); p.add_argument("--extension",action="append"); p.add_argument("--domain"); p.add_argument("--category"); p.add_argument("--subfolder"); p.add_argument("--limit",type=int); p.add_argument("--balanced-sample",action="store_true"); p.add_argument("--force",action="store_true"); a=p.parse_args()
    m=run_import(a.source,a.output,mode=a.mode,extensions=set(a.extension or []),domain=a.domain,category=a.category,subfolder=a.subfolder,limit=a.limit,balanced=a.balanced_sample,force=a.force)
    print(json.dumps({k:m[k] for k in ("mode","examined_count","processed_count","resumed_count","status_counts","text_length_total","page_count_total","paragraph_count_total","table_count_total","duplicate_group_count","error_count","duration_seconds")},ensure_ascii=False,indent=2)); return 0
if __name__=="__main__": raise SystemExit(main())
