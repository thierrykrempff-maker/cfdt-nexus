"""Synthetic-only tests for Protection sociale LOT 1A."""
import json, tempfile, unittest
from pathlib import Path
from unittest import mock
from automation.protection_sociale.document_importer import build_record, duplicate_group_id, extract_docx, extract_pdf, hash_inventory, run_import, sha256_file, stable_document_id

class DocumentImporterTests(unittest.TestCase):
    def setUp(self): self.temp=tempfile.TemporaryDirectory(); self.root=Path(self.temp.name)/"RAW_DOCUMENTS"; self.out=Path(self.temp.name)/"PROCESSED"/"LOT_1A"; self.root.mkdir()
    def tearDown(self): self.temp.cleanup()
    def write(self,relative,data=b"synthetic"):
        p=self.root/relative; p.parent.mkdir(parents=True,exist_ok=True); p.write_bytes(data); return p
    def docx(self,relative="MUTUELLE/NOTICES/synthetic.docx"):
        from docx import Document
        p=self.root/relative; p.parent.mkdir(parents=True,exist_ok=True); d=Document(); d.add_paragraph("Texte entièrement synthétique."); t=d.add_table(rows=1,cols=2); t.cell(0,0).text="A"; t.cell(0,1).text="B"; d.save(p); return p
    def pdf(self,relative="PREVOYANCE/NOTICES/synthetic.pdf"):
        from pypdf import PdfWriter
        p=self.root/relative; p.parent.mkdir(parents=True,exist_ok=True); w=PdfWriter(); w.add_blank_page(width=100,height=100); w.write(str(p)); return p
    def test_ids_and_hashes_stable(self):
        p=self.write("x.pdf"); self.assertEqual(stable_document_id("x.pdf"),stable_document_id("x.pdf")); self.assertEqual(sha256_file(p),sha256_file(p)); self.assertEqual(duplicate_group_id("a"),duplicate_group_id("a"))
    def test_pdf_pages_separator_and_no_ocr(self):
        r=extract_pdf(self.pdf()); self.assertEqual(1,r.page_count); self.assertIn("--- PAGE 1 ---",r.text_content); self.assertIn("probable_image_only_pdf_no_ocr",r.warnings)
    def test_docx_paragraphs_tables(self):
        r=extract_docx(self.docx()); self.assertEqual(1,r.paragraph_count); self.assertEqual(1,r.table_count); self.assertIn("--- TABLE 1 ---",r.text_content)
    def test_empty(self):
        p=self.write("MUTUELLE/empty.pdf",b""); hashes,groups=hash_inventory([p]); r=build_record(p,self.root,hashes[p],groups); self.assertEqual("empty",r.extraction_status); self.assertEqual("",r.text_content); self.assertTrue(r.is_empty_source)
    def test_duplicate_two_records(self):
        a=self.write("MUTUELLE/a.pdf",b"same"); b=self.write("PREVOYANCE/b.pdf",b"same"); hashes,groups=hash_inventory([a,b]); ra=build_record(a,self.root,hashes[a],groups); rb=build_record(b,self.root,hashes[b],groups); self.assertNotEqual(ra.document_id,rb.document_id); self.assertTrue(ra.is_exact_duplicate and rb.is_exact_duplicate); self.assertEqual(ra.duplicate_group_id,rb.duplicate_group_id)
    def test_path_classification_hints(self):
        p=self.docx("MUTUELLE/GARANTIES/tableau_garanties.docx"); hashes,groups=hash_inventory([p]); r=build_record(p,self.root,hashes[p],groups); self.assertEqual("mutuelle",r.document_domain_hint); self.assertEqual("tableau_garanties",r.document_category_hint)
        p=self.docx("PREVOYANCE/NOTICES/notice.docx"); hashes,groups=hash_inventory([p]); self.assertEqual("prévoyance",build_record(p,self.root,hashes[p],groups).document_domain_hint)
    def test_import_outputs_manifest_duplicates_and_relative_paths(self):
        self.docx(); self.write("copy.docx",(self.root/"MUTUELLE/NOTICES/synthetic.docx").read_bytes()); m=run_import(self.root,self.out,mode="import"); self.assertEqual(2,m["processed_count"]); self.assertTrue((self.out/"manifests/import_manifest.json").exists()); self.assertTrue((self.out/"manifests/duplicate_report.json").exists()); self.assertEqual(2,len(list((self.out/"documents").glob("*.json")))); self.assertNotIn(str(self.root),json.dumps(m))
    def test_resume_and_force(self):
        self.docx(); run_import(self.root,self.out,mode="import"); self.assertEqual(1,run_import(self.root,self.out,mode="import")["resumed_count"]); self.assertEqual(1,run_import(self.root,self.out,mode="import",force=True)["processed_count"])
    def test_source_unchanged(self):
        p=self.docx(); before=(p.read_bytes(),p.stat().st_mtime_ns); run_import(self.root,self.out,mode="import"); self.assertEqual(before,(p.read_bytes(),p.stat().st_mtime_ns))
    def test_dry_run_writes_nothing(self):
        self.docx(); m=run_import(self.root,self.out,mode="dry-run"); self.assertEqual(1,m["examined_count"]); self.assertFalse(self.out.exists())
    def test_refuse_output_in_raw(self):
        with self.assertRaises(ValueError): run_import(self.root,self.root/"output",mode="import")
    def test_error_isolated(self):
        p=self.write("broken.pdf",b"not pdf"); self.docx(); m=run_import(self.root,self.out,mode="import"); self.assertEqual(2,m["processed_count"]); self.assertEqual(1,m["error_count"])
    def test_balanced_sample(self):
        self.pdf("a.pdf"); self.pdf("b.pdf"); self.docx("c.docx"); self.write("empty.pdf",b""); self.docx("extra.docx"); m=run_import(self.root,self.out,mode="dry-run",balanced=True); self.assertEqual(4,m["examined_count"]); self.assertEqual(1,m["status_counts"]["empty"])
    def test_no_network_calls(self):
        self.docx()
        with mock.patch("socket.socket",side_effect=AssertionError("network forbidden")): self.assertEqual(1,run_import(self.root,self.out,mode="import")["processed_count"])
if __name__=="__main__": unittest.main()
